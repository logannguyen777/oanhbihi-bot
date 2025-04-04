import os
from sqlalchemy.orm import Session
from database import SessionLocal
from models.web_page import WebPage
from models.document import Document
from models.document_chunk import DocumentChunk
from services.utils.embedding import generate_embedding
from routers.logs_ws import broadcast_log

TEXT_FILE_TYPES = [".txt", ".md"]
PDF_FILE_TYPES = [".pdf"]
DOC_FILE_TYPES = [".doc", ".docx"]
DOWNLOADS_FOLDER = "downloads"

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def split_text(text, max_tokens=500):
    paragraphs = text.split("\\n")
    chunks, current = [], ""
    for para in paragraphs:
        if len((current + para).split()) < max_tokens:
            current += para + "\\n"
        else:
            chunks.append(current.strip())
            current = para + "\\n"
    if current.strip():
        chunks.append(current.strip())
    return chunks

async def train_all():
    await broadcast_log("🚀 Bắt đầu huấn luyện toàn bộ dữ liệu...")
    await train_from_web_pages()
    await train_from_uploaded_files()
    await broadcast_log("✅ Hoàn tất huấn luyện!")

async def train_from_web_pages():
    db: Session = SessionLocal()
    try:
        await broadcast_log("🌐 Huấn luyện từ dữ liệu web đã crawl...")
        pages = db.query(WebPage).filter(WebPage.embedding == None).all()
        count = 0
        for page in pages:
            if not page.content.strip():
                continue
            chunks = split_text(page.content)
            doc = Document(filename=page.url, source="web")
            db.add(doc)
            db.flush()
            for chunk in chunks:
                try:
                    emb = generate_embedding(chunk, db)
                    doc_chunk = DocumentChunk(content=chunk, embedding=emb, document_id=doc.id)
                    db.add(doc_chunk)
                except Exception as e:
                    await broadcast_log(f"❌ Lỗi embedding web chunk: {e}")
            page.embedding = "[trained]"
            count += 1
        db.commit()
        await broadcast_log(f"✅ Đã huấn luyện {count} trang web")
    finally:
        db.close()

async def train_from_uploaded_files():
    db: Session = SessionLocal()
    try:
        await broadcast_log("📄 Huấn luyện từ tài liệu đã upload...")

        if not os.path.exists(DOWNLOADS_FOLDER) or not os.listdir(DOWNLOADS_FOLDER):
            await broadcast_log("⚠️ Chưa có file nào được upload. Vui lòng upload trước khi huấn luyện.")
            return

        files = os.listdir(DOWNLOADS_FOLDER)
        count = 0

        for file in files:
            try:
                await broadcast_log(f"📂 Đang xử lý file: {file}")
                path = os.path.join(DOWNLOADS_FOLDER, file)
                ext = os.path.splitext(file)[1].lower()
                content = ""

                if ext in TEXT_FILE_TYPES:
                    with open(path, "r", encoding="utf-8") as f:
                        content = f.read()
                elif ext in PDF_FILE_TYPES:
                    import fitz
                    doc = fitz.open(path)
                    content = "\n".join([page.get_text() for page in doc])
                elif ext in DOC_FILE_TYPES:
                    import docx
                    doc = docx.Document(path)
                    content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    await broadcast_log(f"⚠️ Không hỗ trợ file: {file}")
                    continue

                if not content.strip():
                    await broadcast_log(f"⚠️ File rỗng hoặc không có nội dung: {file}")
                    continue

                chunks = split_text(content)
                doc = Document(filename=file, source="upload")
                db.add(doc)
                db.flush()

                for chunk in chunks:
                    try:
                        emb = generate_embedding(chunk, db)
                        doc_chunk = DocumentChunk(content=chunk, embedding=emb, document_id=doc.id)
                        db.add(doc_chunk)
                        await broadcast_log(f"🧠 Đã embedding đoạn: {chunk[:50]}...")
                    except Exception as e:
                        await broadcast_log(f"❌ Lỗi embedding đoạn trong {file}: {e}")
                        continue

                await broadcast_log(f"✅ Đã huấn luyện xong: {file}")
                count += 1
                doc.chunk_count = len(chunks)

            except Exception as e:
                await broadcast_log(f"❌ Lỗi xử lý file {file}: {e}")
                continue

        db.commit()
        await broadcast_log(f"🎉 Huấn luyện hoàn tất! Tổng số tài liệu: {count}")

    finally:
        db.close()
